from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level, color="1F3864"):
    p   = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_body(doc, text, bold=False, size=10):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = bold
    run.font.size      = Pt(size)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_after = Pt(3)

def style_table(table, header_color="1F3864", row_alt="EBF0FA"):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold           = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(30, 30, 30)
            if i == 0:
                set_cell_bg(cell, header_color)
            elif i % 2 == 0:
                set_cell_bg(cell, row_alt)

def add_code_block(doc, text):
    for line in text.strip().split('\n'):
        p   = doc.add_paragraph()
        run = p.add_run(line if line else ' ')
        run.font.name       = 'Courier New'
        run.font.size       = Pt(8)
        run.font.color.rgb  = RGBColor(30, 30, 30)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  'F0F4FF')
        p._p.get_or_add_pPr().append(shading)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
run = p.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
run.bold           = True
run.font.size      = Pt(20)
run.font.color.rgb = RGBColor(31, 56, 100)
p.alignment        = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run("Multi-Level Secure File Storage System\nwith Encrypted Vault and Segregated Databases")
run2.font.size      = Pt(13)
run2.font.color.rgb = RGBColor(70, 100, 160)
p2.alignment        = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(30)

# Cover meta table
meta = doc.add_table(rows=5, cols=2)
meta.style = 'Table Grid'
meta_data = [
    ("Project",  "SecureVault – Dual Authentication File Storage System"),
    ("Author",   "Jishan"),
    ("Version",  "1.0"),
    ("Date",     "March 12, 2026"),
    ("Status",   "Development"),
]
for i, (k, v) in enumerate(meta_data):
    meta.cell(i, 0).text = k
    meta.cell(i, 1).text = v
style_table(meta, header_color="1F3864", row_alt="EBF0FA")
# Override: first row is NOT a header here — recolor all rows manually
for i, row in enumerate(meta.rows):
    for j, cell in enumerate(row.cells):
        bg = "D6E4F7" if j == 0 else "F5F9FF"
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.bold = (j == 0)
                run.font.color.rgb = RGBColor(31, 56, 100) if j == 0 else RGBColor(30, 30, 30)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Introduction
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Introduction", 1)

add_heading(doc, "1.1 Purpose", 2)
add_body(doc,
    "This SRS document describes the functional and non-functional requirements of SecureVault, "
    "a multi-level secure file storage system built with Spring Boot. The system demonstrates "
    "real-world cybersecurity concepts including dual-level authentication, AES-256 file encryption, "
    "JWT-based session management, two-factor authentication (2FA), and segregated database storage.")

add_heading(doc, "1.2 Scope", 2)
add_body(doc,
    "SecureVault allows authenticated users to upload, store, and retrieve files across two security "
    "levels. Files stored under Level 1 (normal access) are saved in plaintext, while files stored "
    "under Level 2 (secure access) are encrypted using AES-256 before being persisted to disk. "
    "The system enforces brute-force protection, OTP-based two-factor authentication, and full "
    "security audit logging.")

add_heading(doc, "1.3 Definitions", 2)
defs = [
    ("Level 1",  "Standard access using passwordLevel1; reads/writes normal files"),
    ("Level 2",  "Elevated access using passwordLevel2; reads/writes encrypted files"),
    ("JWT",      "JSON Web Token — stateless session token issued after login"),
    ("OTP",      "One-Time Password sent via email for 2FA verification"),
    ("AES-256",  "Advanced Encryption Standard with 256-bit key"),
    ("BCrypt",   "Password hashing algorithm used to store credentials securely"),
]
t = doc.add_table(rows=len(defs)+1, cols=2)
t.style = 'Table Grid'
t.cell(0,0).text = "Term"
t.cell(0,1).text = "Meaning"
for i,(k,v) in enumerate(defs, start=1):
    t.cell(i,0).text = k
    t.cell(i,1).text = v
style_table(t)

doc.add_paragraph()
add_heading(doc, "1.4 Technology Stack", 2)
stack = [
    ("Layer",          "Technology"),
    ("Backend",        "Java 17, Spring Boot 3.x"),
    ("Security",       "Spring Security, JWT (JJWT)"),
    ("Database",       "MySQL (securevault_db)"),
    ("Encryption",     "AES-256 (Java Cipher API)"),
    ("Password Hash",  "BCrypt"),
    ("2FA / Email",    "JavaMail (Gmail SMTP)"),
    ("Frontend",       "HTML, CSS, JavaScript (static)"),
]
ts = doc.add_table(rows=len(stack), cols=2)
ts.style = 'Table Grid'
for i,(k,v) in enumerate(stack):
    ts.cell(i,0).text = k
    ts.cell(i,1).text = v
style_table(ts)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Overall Description
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "2. Overall Description", 1)

add_heading(doc, "2.1 System Overview", 2)
arch = """
User (Browser)
     │
     ▼
┌─────────────────────────────────────────────┐
│              Spring Boot Backend             │
│  ┌──────────┐  ┌──────────┐  ┌───────────┐  │
│  │AuthCtrl  │  │FileCtrl  │  │JwtFilter  │  │
│  └────┬─────┘  └────┬─────┘  └─────┬─────┘  │
│       │             │              │         │
│  ┌────▼─────┐  ┌────▼─────┐  ┌────▼──────┐  │
│  │AuthSvc   │  │FileSvc   │  │JwtUtil    │  │
│  └────┬─────┘  └────┬─────┘  └───────────┘  │
│       │             │                        │
│  ┌────▼─────┐  ┌────▼──────────────────────┐ │
│  │ MySQL DB │  │  /storage/normal/          │ │
│  │ users    │  │  /storage/secure/ (AES-256)│ │
│  │ files    │  └───────────────────────────┘ │
│  └──────────┘                                │
└─────────────────────────────────────────────┘
"""
add_code_block(doc, arch)

doc.add_paragraph()
add_heading(doc, "2.2 User Roles", 2)
roles = [
    ("Role",         "Description"),
    ("Guest",        "Can register a new account"),
    ("Level 1 User", "Logged in with Password 1 — accesses normal file vault"),
    ("Level 2 User", "Logged in with Password 2 — accesses encrypted vault"),
]
tr = doc.add_table(rows=len(roles), cols=2)
tr.style = 'Table Grid'
for i,(k,v) in enumerate(roles):
    tr.cell(i,0).text = k
    tr.cell(i,1).text = v
style_table(tr)

doc.add_paragraph()
add_heading(doc, "2.3 Assumptions & Constraints", 2)
for item in [
    "MySQL must be running locally on port 3306.",
    "Gmail SMTP credentials must be configured for 2FA OTP delivery.",
    "File upload size is limited to 50 MB per file.",
    "JWT tokens expire after 1 hour.",
    "OTP is valid for 5 minutes.",
]:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Functional Requirements
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "3. Functional Requirements", 1)

fr = [
    ("FR-01 — User Registration", [
        "User provides: username, passwordLevel1, passwordLevel2, email, securityQuestion, securityAnswer.",
        "Both passwords are hashed with BCrypt before storage.",
        "Security answer is also BCrypt-hashed.",
    ]),
    ("FR-02 — Dual-Level Authentication", [
        "Level 1 Login: validates passwordLevel1 → issues JWT with accessLevel = LEVEL1.",
        "Level 2 Login: validates passwordLevel2 → issues JWT with accessLevel = LEVEL2.",
        "If 2FA is enabled, an OTP is emailed; user must submit OTP to receive the final JWT.",
    ]),
    ("FR-03 — Two-Factor Authentication (2FA)", [
        "User can enable/disable 2FA from the dashboard.",
        "On login, if 2FA is enabled, a 6-digit OTP is sent to the registered email.",
        "OTP expires after 5 minutes. Failed OTP resets the flow.",
    ]),
    ("FR-04 — Brute Force Protection", [
        "After 5 consecutive failed login attempts, the account is locked.",
        "Locked accounts cannot log in until unlocked via the password recovery flow.",
    ]),
    ("FR-05 — Password Recovery", [
        "User answers their registered security question.",
        "If correct, they may set a new passwordLevel1 and passwordLevel2.",
    ]),
    ("FR-06 — File Upload", [
        "Authenticated users upload files via the dashboard.",
        "Level 1: file is stored as-is under /storage/normal/.",
        "Level 2: file is AES-256 encrypted before being saved under /storage/secure/.",
        "File metadata (name, owner, path, access level, upload time) is stored in MySQL.",
    ]),
    ("FR-07 — File Download", [
        "Users can list and download their own files.",
        "Level 2 files are decrypted on-the-fly before being returned to the browser.",
        "Users cannot access files belonging to other users.",
    ]),
    ("FR-08 — File Delete", [
        "Users can delete their own files.",
        "Deletes both the physical file from disk and the metadata from the database.",
    ]),
    ("FR-09 — JWT-Based Session Management", [
        "Every protected API endpoint requires a valid Authorization: Bearer <token> header.",
        "JwtFilter intercepts every request, validates the token, and injects the security context.",
    ]),
]
for title, bullets in fr:
    add_heading(doc, title, 2)
    for b in bullets:
        add_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Non-Functional Requirements
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "4. Non-Functional Requirements", 1)
nfr = [
    ("ID",     "Category",     "Requirement"),
    ("NFR-01", "Security",     "All passwords stored as BCrypt hashes; never stored in plaintext"),
    ("NFR-02", "Security",     "Sensitive files encrypted with AES-256 at rest"),
    ("NFR-03", "Security",     "JWT signed with HMAC-SHA256; expiry enforced server-side"),
    ("NFR-04", "Availability", "Application runs on localhost:8080; accessible via any browser"),
    ("NFR-05", "Performance",  "File upload/download under 3 seconds for files up to 10 MB"),
    ("NFR-06", "Scalability",  "MySQL schema auto-updated via Hibernate DDL"),
    ("NFR-07", "Usability",    "Single-page HTML/CSS/JS frontend; no framework dependency"),
    ("NFR-08", "Auditability", "Every login, upload, download, and failure logged in application log"),
]
tn = doc.add_table(rows=len(nfr), cols=3)
tn.style = 'Table Grid'
for i, row_data in enumerate(nfr):
    for j, val in enumerate(row_data):
        tn.cell(i, j).text = val
style_table(tn)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — System Components
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "5. System Components", 1)
comps = [
    ("Component",            "Class / File",                "Responsibility"),
    ("Entry Point",          "SecurevaultApplication.java", "Boots Spring context"),
    ("Auth Controller",      "AuthController.java",         "Register, login, OTP, recovery endpoints"),
    ("File Controller",      "FileController.java",         "Upload, download, list, delete endpoints"),
    ("Auth Service",         "AuthService.java",            "Business logic for auth & account management"),
    ("File Service",         "FileService.java",            "File I/O, storage path routing"),
    ("Encryption Service",   "EncryptionService.java",      "AES-256 encrypt / decrypt"),
    ("Two-Factor Service",   "TwoFactorService.java",       "OTP generation & email dispatch"),
    ("JWT Utility",          "JwtUtil.java",                "Token generation, validation, claim parsing"),
    ("JWT Filter",           "JwtFilter.java",              "Intercepts HTTP requests, validates tokens"),
    ("Security Config",      "SecurityConfig.java",         "Spring Security rules, CORS, public routes"),
    ("DataSource Config",    "DataSourceConfig.java",       "MySQL datasource configuration"),
    ("User Entity",          "User.java",                   "DB model: credentials, 2FA, lock status"),
    ("FileMetadata Entity",  "FileMetadata.java",           "DB model: file name, owner, level, path"),
]
tc2 = doc.add_table(rows=len(comps), cols=3)
tc2.style = 'Table Grid'
for i, row_data in enumerate(comps):
    for j, val in enumerate(row_data):
        tc2.cell(i, j).text = val
style_table(tc2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Use Case Summary
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "6. Use Case Summary", 1)
uc = """
┌─────────────────────────────────────────────────────┐
│                      SecureVault                    │
│                                                     │
│  [User] ────► Register                              │
│  [User] ────► Login (Level 1 / Level 2)             │
│       └──────► [If 2FA] Submit OTP                  │
│  [Level1 User] ► Upload / Download / Delete Normal  │
│  [Level2 User] ► Upload / Download / Delete Encrypted│
│  [User] ────► Enable / Disable 2FA                  │
│  [User] ────► Recover Password via Security Question │
└─────────────────────────────────────────────────────┘
"""
add_code_block(doc, uc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Data Flow
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading(doc, "7. Data Flow — Secure File Upload (Level 2)", 1)
df = """
Browser ──[POST /api/files/upload + JWT(LEVEL2)]──► FileController
                                                         │
                                                    FileService
                                                         │
                                          EncryptionService (AES-256)
                                                         │
                                 /storage/secure/<timestamp>_<filename>
                                                         │
                                          FileMetadataRepository
                                                         │
                                       MySQL → file_metadata table
"""
add_code_block(doc, df)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("End of SRS Document — SecureVault v1.0")
run.italic          = True
run.font.size       = Pt(9)
run.font.color.rgb  = RGBColor(120, 120, 120)
p.alignment         = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"C:\Users\mdjis\Desktop\DualPasswordFileSystem\securevault\SRS_Report.docx"
doc.save(out)
print(f"Saved: {out}")

